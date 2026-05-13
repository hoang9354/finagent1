"""
modules/report_docx.py — Word Document Report Generator (v3)
=============================================================
Produces a polished .docx investment report with:
  - Professional cover page (navy banner)
  - Running header (report title + date) on every body page
  - Correct "Page X of Y" footer
  - Asset summary statistics table
  - All 5 chart images embedded (one per page, labelled)
  - Full 7-section AI analysis body with styled headings
"""

import os
import re
from datetime import datetime
from typing import Dict, List, Optional

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

# ─── Palette ──────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1F, 0x35, 0x64)
BLUE  = RGBColor(0x2E, 0x75, 0xB6)
LGREY = RGBColor(0xF2, 0xF7, 0xFD)
GREY  = RGBColor(0x26, 0x26, 0x26)
MID   = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x1E, 0x8B, 0x4C)
RED   = RGBColor(0xC0, 0x00, 0x00)
GOLD  = RGBColor(0xC9, 0xA0, 0x2E)

# Chart labels matching the filenames
CHART_LABELS = {
    "chart1_price_volume.png":        "Chart 1 — Price Trend & Volume",
    "chart2_correlation_heatmap.png": "Chart 2 — Daily Return Correlation Heatmap",
    "chart3_return_distribution.png": "Chart 3 — Return Distribution (Histogram + KDE)",
    "chart4_rolling_stats.png":       "Chart 4 — Rolling Statistics & Bollinger Bands",
    "chart5_cumulative_returns.png":  "Chart 5 — Cumulative Returns Comparison (Bonus)",
}


# ─── XML helpers ──────────────────────────────────────────────────────────────

def _cell_shade(cell, hex_fill: str):
    """Set cell background via OOXML — works in all Word versions."""
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn("w:shd")):
        tcPr.remove(e)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill.upper())
    tcPr.append(shd)


def _bottom_border(para, color: str = "2E75B6", size: int = 10):
    """Draw a coloured bottom border line under a paragraph."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color.upper())
    pBdr.append(bot)
    pPr.append(pBdr)


def _left_bar_para(para, color_hex: str = "2E75B6", width_twips: int = 40):
    """Add a thick left border to a paragraph (pull-quote / story style)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(width_twips))
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), color_hex.upper())
    pBdr.append(left)
    pPr.append(pBdr)


def _insert_field(run, instruction: str):
    """
    Insert a Word field (PAGE or NUMPAGES) into a run.
    Correct structure: fldChar(begin) → instrText → fldChar(end).
    """
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fc_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run._r.append(instr)

    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    run._r.append(fc_end)


def _styled_run(para, text: str = "", size: float = 9,
                bold: bool = False, color: RGBColor = None) -> object:
    r = para.add_run(text)
    r.font.name  = "Arial"
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.color.rgb = color or MID
    return r


# ─── Header & Footer ──────────────────────────────────────────────────────────

def _setup_header_footer(section, title: str, generated: str):
    """
    Header  : report title (left) + date (right)
    Footer  : thin top border + centred "Page X of Y"
    Applied only to body sections (not cover).
    """
    # ── Header ────────────────────────────────────────────────────────────
    header = section.header
    hpara  = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hpara.clear()
    hpara.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Left: report title
    rl = hpara.add_run(title)
    rl.font.name  = "Arial"
    rl.font.size  = Pt(9)
    rl.font.bold  = True
    rl.font.color.rgb = NAVY

    # Tab to right-align date
    from docx.oxml import OxmlElement as OE
    tab = OE("w:tab"); hpara._p.append(tab)

    # Set a right tab stop at page width
    pPr  = hpara._p.get_or_add_pPr()
    tabs = OE("w:tabs")
    tab_stop = OE("w:tab")
    tab_stop.set(qn("w:val"), "right")
    tab_stop.set(qn("w:pos"), "9350")   # ~6.5 inches in twips
    tabs.append(tab_stop)
    pPr.append(tabs)

    rd = hpara.add_run(f"\t{generated}")
    rd.font.name  = "Arial"
    rd.font.size  = Pt(9)
    rd.font.color.rgb = MID

    # Thin blue bottom border on header paragraph
    _bottom_border(hpara, "2E75B6", 6)

    # ── Footer ────────────────────────────────────────────────────────────
    footer = section.footer
    fpara  = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fpara.clear()
    fpara.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Thin top border
    _top_border_para(fpara, "2E75B6", 6)

    _styled_run(fpara, "FinAgent Investment Report   |   Page ", 9, False, MID)
    _insert_field(_styled_run(fpara, size=9, color=MID), "PAGE")
    _styled_run(fpara, " of ", 9, False, MID)
    _insert_field(_styled_run(fpara, size=9, color=MID), "NUMPAGES")


def _top_border_para(para, color: str = "2E75B6", size: int = 6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    str(size))
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color.upper())
    pBdr.append(top)
    pPr.append(pBdr)


# ─── Cover Page ───────────────────────────────────────────────────────────────

def _cover(doc: Document, tickers: List[str], generated: str):
    """
    Full-width navy banner with white title, then meta block below.
    Cover uses its own section so the header/footer is suppressed.
    """
    # Navy banner table
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    _cell_shade(cell, "1F3564")

    for txt, sz, bold, clr in [
        ("",                         5,  False, WHITE),
        ("FinAgent",                42,  True,  WHITE),
        ("Investment Analysis Report", 20, False, RGBColor(0xBD, 0xD7, 0xEE)),
        ("AI-Powered Financial Data Pipeline", 12, False, RGBColor(0x9D, 0xC3, 0xE6)),
        ("",                         5,  False, WHITE),
    ]:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.name = "Arial"; r.font.size = Pt(sz)
        r.font.bold = bold;    r.font.color.rgb = clr

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(12)

    # Horizontal rule
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after  = Pt(14)
    _bottom_border(rule, "2E75B6", 16)

    # Meta info block
    for label, value in [
        ("Assets Analysed", "   |   ".join(tickers)),
        ("Generated",        generated),
        ("AI Provider",      "Groq — llama-3.3-70b-versatile"),
        ("Data Source",      "Yahoo Finance (yfinance)"),
    ]:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_after = Pt(6)
        rl = mp.add_run(f"{label}:  ")
        rl.font.name = "Arial"; rl.font.size = Pt(11)
        rl.font.bold = True;    rl.font.color.rgb = NAVY
        rv = mp.add_run(value)
        rv.font.name = "Arial"; rv.font.size = Pt(11)
        rv.font.color.rgb = GREY

    # Disclaimer at the bottom
    sp2 = doc.add_paragraph()
    sp2.paragraph_format.space_after = Pt(30)

    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = disc.add_run(
        "This report is generated for academic purposes only and does not constitute "
        "financial advice. All analysis is based on historical data."
    )
    rd.font.name   = "Arial"
    rd.font.size   = Pt(8)
    rd.font.italic = True
    rd.font.color.rgb = MID

    doc.add_page_break()


# ─── Stats Table ──────────────────────────────────────────────────────────────

def _stats_table(doc: Document, cleaned: Dict, company_info: Optional[Dict]):
    # Section heading
    sh = doc.add_paragraph()
    sh.paragraph_format.space_before = Pt(0)
    sh.paragraph_format.space_after  = Pt(6)
    rsh = sh.add_run("Asset Summary Statistics")
    rsh.font.name = "Arial"; rsh.font.size = Pt(17)
    rsh.font.bold = True;    rsh.font.color.rgb = NAVY
    _bottom_border(sh, "1F3564", 14)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    ri = intro.add_run(
        "Key price and risk metrics for all tracked assets over the past 12 months."
    )
    ri.font.name = "Arial"; ri.font.size = Pt(10); ri.font.color.rgb = MID

    COLS = [
        ("Ticker",     0.62), ("Company",   1.50), ("Price",    0.72),
        ("52w High",   0.72), ("52w Low",   0.72), ("Period Ret", 0.82),
        ("Ann. Vol",   0.72), ("Sharpe",    0.62), ("Trend",    1.02),
    ]

    tbl = doc.add_table(rows=1, cols=len(COLS))
    tbl.style = "Table Grid"

    # Header row
    for i, (label, w) in enumerate(COLS):
        c = tbl.rows[0].cells[i]
        c.width = Inches(w)
        _cell_shade(c, "1F3564")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Twips(60)
        p.paragraph_format.space_after  = Twips(60)
        r = p.add_run(label)
        r.font.name = "Arial"; r.font.size = Pt(9)
        r.font.bold = True;    r.font.color.rgb = WHITE

    # Data rows
    for idx, (ticker, df) in enumerate(cleaned.items()):
        close  = df["Close"].dropna()
        ret    = df["Daily_Return"].dropna()
        price  = close.iloc[-1]
        hi, lo = close.max(), close.min()
        p_ret  = ((close.iloc[-1] / close.iloc[0]) - 1) * 100
        ann_v  = ret.std() * (252 ** 0.5) * 100
        sharpe = (ret.mean() / ret.std() * np.sqrt(252)) if ret.std() > 0 else 0
        ma30   = df["MA_30"].iloc[-1] if "MA_30" in df.columns else None
        trend  = ""
        if ma30 is not None and not pd.isna(ma30):
            trend = "↑ Above MA30" if close.iloc[-5:].mean() > ma30 else "↓ Below MA30"
        name  = ((company_info or {}).get(ticker, {}).get("longName") or ticker)[:22]
        vals  = [ticker, name, f"${price:.2f}", f"${hi:.2f}", f"${lo:.2f}",
                 f"{p_ret:+.2f}%", f"{ann_v:.2f}%", f"{sharpe:.2f}", trend]
        bg    = "EBF3FB" if idx % 2 == 0 else "FFFFFF"
        row   = tbl.add_row()

        for j, (val, (_, w)) in enumerate(zip(vals, COLS)):
            c = row.cells[j]
            c.width = Inches(w)
            _cell_shade(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Twips(55)
            p.paragraph_format.space_after  = Twips(55)
            r = p.add_run(str(val))
            r.font.name = "Arial"; r.font.size = Pt(9)
            if j == 0:
                r.font.bold = True; r.font.color.rgb = NAVY
            elif j == 5:
                r.font.bold = True
                r.font.color.rgb = GREEN if p_ret >= 0 else RED
            elif j == 8:
                r.font.color.rgb = GREEN if "Above" in str(val) else (RED if "Below" in str(val) else GREY)
            else:
                r.font.color.rgb = GREY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    doc.add_page_break()


# ─── Charts Section ───────────────────────────────────────────────────────────

def _insert_charts(doc: Document, charts_dir: str):
    """
    Embed all 5 chart PNGs into the document.
    Each chart gets a labelled heading, the image, and a caption.
    """
    if not charts_dir or not os.path.isdir(charts_dir):
        return

    # Section heading
    sh = doc.add_paragraph()
    sh.paragraph_format.space_before = Pt(0)
    sh.paragraph_format.space_after  = Pt(6)
    rsh = sh.add_run("Data Visualisations")
    rsh.font.name = "Arial"; rsh.font.size = Pt(17)
    rsh.font.bold = True;    rsh.font.color.rgb = NAVY
    _bottom_border(sh, "1F3564", 14)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    ri = intro.add_run(
        "All charts are generated from cleaned historical data. "
        "Moving averages and Bollinger Bands use 7-day, 30-day, and 20-day windows respectively."
    )
    ri.font.name = "Arial"; ri.font.size = Pt(10); ri.font.color.rgb = MID

    captions = {
        "chart1_price_volume.png":
            "Close price with 7-day and 30-day moving averages alongside daily trading volume.",
        "chart2_correlation_heatmap.png":
            "Pearson correlation coefficients between daily returns. Values near +1 indicate assets move together; near 0 indicates independence.",
        "chart3_return_distribution.png":
            "Histogram and KDE of daily returns with mean (dashed) and ±1σ (dotted) markers. Skew and kurtosis values indicate tail risk.",
        "chart4_rolling_stats.png":
            "Close price overlaid with MA7, MA30, and Bollinger Bands (20-day, ±2σ). Price touching the upper band may signal overbought conditions.",
        "chart5_cumulative_returns.png":
            "Normalised buy-and-hold cumulative return from the start of the analysis period. Allows direct performance comparison across assets.",
    }

    for filename, label in CHART_LABELS.items():
        img_path = os.path.join(charts_dir, filename)
        if not os.path.isfile(img_path):
            continue

        # Chart label heading
        ch = doc.add_paragraph()
        ch.paragraph_format.space_before = Pt(14)
        ch.paragraph_format.space_after  = Pt(4)
        rch = ch.add_run(label)
        rch.font.name = "Arial"; rch.font.size = Pt(13)
        rch.font.bold = True;    rch.font.color.rgb = BLUE

        # Image — fit to page width with margins
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_para.paragraph_format.space_after = Pt(4)
        run = img_para.add_run()
        run.add_picture(img_path, width=Inches(6.3))

        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(16)
        rc = cap.add_run(captions.get(filename, ""))
        rc.font.name   = "Arial"
        rc.font.size   = Pt(9)
        rc.font.italic = True
        rc.font.color.rgb = MID

    doc.add_page_break()


# ─── Inline text formatter ────────────────────────────────────────────────────

def _inline_runs(para, text: str, size: Pt = Pt(11)):
    """Parse **bold** and *italic* markers and add styled runs."""
    pos = 0
    for m in re.finditer(r"(\*\*(.+?)\*\*|\*(.+?)\*)", text):
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            r.font.name = "Arial"; r.font.size = size; r.font.color.rgb = GREY
        inner = m.group(2) or m.group(3)
        bold  = m.group(0).startswith("**")
        r = para.add_run(inner)
        r.font.name = "Arial"; r.font.size = size
        r.font.bold = bold;    r.font.italic = not bold
        r.font.color.rgb = NAVY if bold else MID
        pos = m.end()
    if pos < len(text):
        r = para.add_run(text[pos:])
        r.font.name = "Arial"; r.font.size = size; r.font.color.rgb = GREY


# ─── Heading renderer ─────────────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(15)
        r.font.bold = True;    r.font.color.rgb = NAVY
        _bottom_border(p, "2E75B6", 8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(13)
        r.font.bold = True;    r.font.color.rgb = BLUE
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(11)
        r.font.bold = True;    r.font.color.rgb = MID


# ─── Markdown → DOCX parser ──────────────────────────────────────────────────


def _render_stock_stories(doc: Document, md: str):
    """
    Detect the "Stock Stories & Key Drivers" section in the markdown and
    render it with a distinctive pull-quote style:
      - Navy section heading
      - Per-stock sub-heading with coloured left accent bar
      - Story paragraphs indented with a light background feel
    Falls back to normal rendering if the section is not found.
    """
    # Find the section boundaries
    lines = md.splitlines()
    in_stories = False
    stories_lines = []
    after_lines   = []

    for line in lines:
        stripped = line.strip()
        # Detect start of Stock Stories section (## 1. Stock Stories...)
        if re.match(r"^#{1,2}\s+1\.\s+Stock Stories", stripped, re.IGNORECASE):
            in_stories = True
            continue
        # Detect start of any subsequent ## section — end of stories
        if in_stories and re.match(r"^#{1,2}\s+[2-9]", stripped):
            in_stories = False
            after_lines.append(line)
            continue
        if in_stories:
            stories_lines.append(line)
        else:
            after_lines.append(line)

    if not stories_lines:
        # Section not found — render everything normally
        _parse_markdown(doc, md)
        return

    # ── Render the Stock Stories heading ──────────────────────────────────
    sh = doc.add_paragraph()
    sh.paragraph_format.space_before = Pt(0)
    sh.paragraph_format.space_after  = Pt(6)
    rsh = sh.add_run("1. Stock Stories & Key Drivers")
    rsh.font.name = "Arial"; rsh.font.size = Pt(15)
    rsh.font.bold = True;    rsh.font.color.rgb = NAVY
    _bottom_border(sh, "2E75B6", 8)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    ri = intro.add_run(
        "A plain-English narrative of what drove each stock over the analysis period — "
        "covering the business story, the price story, and the news catalyst story."
    )
    ri.font.name = "Arial"; ri.font.size = Pt(10); ri.font.italic = True
    ri.font.color.rgb = MID

    # Colour cycle for per-stock accent bars
    ACCENT_COLORS = ["1F3564", "2E75B6", "00703C", "C9A02E"]
    stock_idx = 0

    i = 0
    while i < len(stories_lines):
        line = stories_lines[i].rstrip()

        # Skip blanks between stocks
        if not line.strip():
            i += 1
            continue

        # Per-stock sub-heading: ### TICKER or **TICKER**
        m_stock = re.match(r"^#{2,4}\s+(.+)$", line) or re.match(r"^\*\*(.+)\*\*\s*$", line)
        if m_stock:
            accent = ACCENT_COLORS[stock_idx % len(ACCENT_COLORS)]
            stock_idx += 1

            # Spacer before each stock block
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(4)

            # Stock name heading with left accent bar
            sh2 = doc.add_paragraph()
            sh2.paragraph_format.space_before = Pt(4)
            sh2.paragraph_format.space_after  = Pt(6)
            sh2.paragraph_format.left_indent  = Inches(0.15)
            _left_bar_para(sh2, accent, 36)
            rs2 = sh2.add_run(m_stock.group(1).strip())
            rs2.font.name = "Arial"; rs2.font.size = Pt(13)
            rs2.font.bold = True
            rs2.font.color.rgb = RGBColor(
                int(accent[0:2], 16),
                int(accent[2:4], 16),
                int(accent[4:6], 16),
            )
            i += 1
            continue

        # Paragraph label lines: **Paragraph N — ...**
        m_label = re.match(r"^\*\*Paragraph\s+\d+\s*[—–-]\s*(.+?)\*\*", line)
        if m_label:
            lp = doc.add_paragraph()
            lp.paragraph_format.space_before = Pt(8)
            lp.paragraph_format.space_after  = Pt(2)
            lp.paragraph_format.left_indent  = Inches(0.25)
            rl = lp.add_run(m_label.group(1).strip())
            rl.font.name = "Arial"; rl.font.size = Pt(10)
            rl.font.bold = True;    rl.font.color.rgb = BLUE
            i += 1
            continue

        # Bullet inside story
        mb = re.match(r"^(\s*)[-*+] (.+)$", line)
        if mb:
            lvl = len(mb.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Inches(0.45 + 0.2 * lvl)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(3)
            _inline_runs(p, mb.group(2), Pt(10.5))
            i += 1
            continue

        # Story body paragraph — indented with left accent bar in light blue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.1)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.line_spacing = Pt(14)
        _inline_runs(p, line)
        i += 1

    # Divider before rest of report
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(14)
    div.paragraph_format.space_after  = Pt(14)
    _bottom_border(div, "2E75B6", 6)

    # ── Render remaining sections normally ────────────────────────────────
    _parse_markdown(doc, "\n".join(after_lines))

def _parse_markdown(doc: Document, md: str):
    """
    Convert Groq markdown output into styled Word paragraphs.
    Handles: # ## (H1), ### (H2), #### (H3), bullets, numbered lists, body text.
    Suppresses consecutive blank lines (max one spacer).
    """
    blanks = 0
    for raw in md.splitlines():
        line = raw.rstrip()

        # Blank line
        if not line.strip():
            blanks += 1
            if blanks == 1:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(2)
            continue
        blanks = 0

        # Skip horizontal rules
        if re.match(r"^(-{3,}|\*{3,}|_{3,})$", line.strip()):
            continue

        # H1: # or ##
        if re.match(r"^#{1,2} ", line):
            _heading(doc, re.sub(r"^#{1,2} ", "", line).strip(), 1)
            continue

        # H2: ###
        if re.match(r"^### ", line):
            _heading(doc, re.sub(r"^### ", "", line).strip(), 2)
            continue

        # H3: ####
        if re.match(r"^#### ", line):
            _heading(doc, re.sub(r"^#### ", "", line).strip(), 3)
            continue

        # Bullet point
        mb = re.match(r"^(\s*)[-*+] (.+)$", line)
        if mb:
            lvl = len(mb.group(1)) // 2
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Inches(0.3 + 0.2 * lvl)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(3)
            _inline_runs(p, mb.group(2), Pt(10.5))
            continue

        # Numbered list
        mn = re.match(r"^\d+\. (.+)$", line)
        if mn:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(3)
            _inline_runs(p, mn.group(1), Pt(10.5))
            continue

        # Body paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.line_spacing = Pt(14)
        _inline_runs(p, line)


# ─── Public API ───────────────────────────────────────────────────────────────

def save_report_as_docx(
    analysis_text: str,
    tickers: List[str],
    cleaned: Dict,
    company_info: Optional[Dict] = None,
    output_dir: str = "outputs/reports",
    charts_dir: Optional[str] = None,
) -> str:
    """
    Build and save a formatted Word report.

    Parameters
    ----------
    analysis_text : Markdown text from Groq.
    tickers       : List of ticker symbols.
    cleaned       : Dict of cleaned DataFrames (for stats table).
    company_info  : Optional company profile dicts.
    output_dir    : Where to save the .docx file.
    charts_dir    : Path to the run's chart folder. If provided, all
                    5 PNGs are embedded in the report.

    Returns
    -------
    Full path to the saved .docx file.
    """
    os.makedirs(output_dir, exist_ok=True)
    generated = datetime.now().strftime("%Y-%m-%d  %H:%M")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path  = os.path.join(output_dir, f"FinAgent_Report_{timestamp}.docx")
    title     = "FinAgent Investment Analysis Report"

    doc = Document()

    # ── Page setup: A4, 1-inch margins ────────────────────────────────────
    for sec in doc.sections:
        sec.page_width    = Twips(11906)
        sec.page_height   = Twips(16838)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)

    # ── Document default style ────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name  = "Arial"
    normal.font.size  = Pt(11)
    normal.font.color.rgb = GREY
    normal.paragraph_format.space_after  = Pt(5)
    normal.paragraph_format.space_before = Pt(0)

    # ── Cover page (section 1 — no header/footer) ─────────────────────────
    _cover(doc, tickers, generated)

    # ── Add a new section for body pages so we can add header/footer ──────
    # The page break from _cover already ended the first section.
    # We configure the current (only) section's header/footer here.
    # Word applies section 0 settings to all pages unless a section break
    # is inserted; since python-docx creates one section by default we
    # apply header/footer to it. The cover page suppresses header/footer
    # by virtue of "different first page" setting.
    section = doc.sections[0]
    section.different_first_page_header_footer = True   # cover has no header
    _setup_header_footer(section, title, generated)

    # ── Stats table ───────────────────────────────────────────────────────
    _stats_table(doc, cleaned, company_info)

    # ── Charts ────────────────────────────────────────────────────────────
    if charts_dir:
        _insert_charts(doc, charts_dir)

    # ── AI analysis body ──────────────────────────────────────────────────
    _render_stock_stories(doc, analysis_text)

    doc.save(out_path)
    return out_path
